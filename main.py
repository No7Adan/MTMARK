import os
from langchain_openai import ChatOpenAI
from nltk.translate.bleu_score import sentence_bleu, SmoothingFunction
from langchain_core.pydantic_v1 import BaseModel, Field
from collections import Counter
from docx import Document
from docx.shared import Inches
from docx.oxml.ns import qn
from docx.shared import Pt
import pandas as pd
import requests
import random
import html
from hashlib import md5
from utils.AuthV3Util import addAuthParams
import matplotlib.pyplot as plt
import matplotlib
import tkinter as tk
from tkinter import messagebox
import numpy as np


# 创建Word文档和预定义字体
doc = Document()
doc.styles['Normal'].font.name = 'Times New Roman'
doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
# 为TK可视化界面定义了输入参数
def preparation():
    def submit():
        # 从 seg_number_entry 获取值
        global seg_number, pair
        seg_number = seg_number_entry.get()
        pair_option = pair_var.get()
        pair = "中 --> En" if pair_option == 1 else "En --> 中"
        root.destroy()
        main(seg_number, pair)

    # 定义并运行TK可视化程序
    root = tk.Tk()
    root.title("Preparation")
    root.geometry("300x200")
    root.resizable(True, True)

    padding = {'padx': 10, 'pady': 10}

    tk.Label(root, text="The Number of Segments").grid(row=0, **padding)
    seg_number_entry = tk.Entry(root)
    seg_number_entry.grid(row=0, column=1, sticky="ew", **padding)
    pair_var = tk.IntVar()
    tk.Radiobutton(root, text="中 --> En", variable=pair_var, value=1).grid(row=4, column=0, **padding)
    tk.Radiobutton(root, text="En --> 中", variable=pair_var, value=2).grid(row=4, column=1, **padding)

    tk.Button(root, text='Submit', command=submit).grid(row=5, column=0, columnspan=2, **padding)
    root.grid_columnconfigure(1, weight=1)
    root.mainloop()

def get_user_input(pair):
    def submit():
        global refer_text, source_text
        refer_text = refer_text_entry.get("1.0", "end-1c")
        source_text = source_text_entry.get("1.0", "end-1c")
        root.destroy()

# 定义并运行TK可视化程序
    root = tk.Tk()
    root.title("Translation Input")
    root.geometry("600x400")
    root.resizable(True, True)

    padding = {'padx': 10, 'pady': 10}

    tk.Label(root, text="Refer Text").grid(row=1, **padding)
    tk.Label(root, text="Source Text").grid(row=2, **padding)

    refer_text_entry = tk.Text(root, height=5, width=50)
    source_text_entry = tk.Text(root, height=5, width=50)

    refer_text_entry.grid(row=1, column=1, sticky="ew", **padding)
    source_text_entry.grid(row=2, column=1, sticky="ew", **padding)

    tk.Button(root, text='Submit', command=submit).grid(row=5, column=0, columnspan=2, **padding)

    root.grid_columnconfigure(1, weight=1)
    root.grid_rowconfigure(1, weight=1)
    root.grid_rowconfigure(2, weight=1)

    root.mainloop()

# 定义颜色，为下文输出增加色彩
class Colors:
    HEADER = '\033[95m'
    OKBLUE = '\033[94m'
    OKCYAN = '\033[96m'
    OKGREEN = '\033[92m'
    WARNING = '\033[93m'
    FAIL = '\033[91m'
    ENDC = '\033[0m'
    BOLD = '\033[1m'
    UNDERLINE = '\033[4m'

# 定义机器翻译API运行和质量测评主程序运行的函数
total_result = []
def run_translation(refer_text, source_text, pair):
    MT_texts = []
    single_result = []
    if pair == "中 --> En":
        source_lang = 'zh'
        target_lang = 'en'
    elif pair == "En --> 中":
        source_lang = 'en'
        target_lang = 'zh'
    else:
        raise ValueError("Invalid pair value. Expected '中 --> En' or 'En --> 中'.") 

    # 百度翻译API
    def BMT(source_text, pair):
        appid = 'APIKEY_01657528'
        appkey = 'APIKEY_2BNI1bea'
        endpoint = 'http://api.fanyi.baidu.com'
        path = '/api/trans/vip/translate'
        url = endpoint + path

        if pair == "中 --> En":
            source_lang = 'zh'
            target_lang = 'en'
        elif pair == "En --> 中":
            source_lang = 'en'
            target_lang = 'zh'
        else:
            raise ValueError("Invalid pair value. Expected '中 --> En' or 'En --> 中'.")

        def make_md5(s, encoding='utf-8'):
            return md5(s.encode(encoding)).hexdigest()

        salt = random.randint(32768, 65536)
        sign = make_md5(appid + source_text + str(salt) + appkey)

        headers = {'Content-Type': 'application/x-www-form-urlencoded'}
        payload = {'appid': appid, 'q': source_text, 'from': source_lang, 'to': target_lang, 'salt': salt, 'sign': sign}

        r = requests.post(url, params=payload, headers=headers)
        result = r.json()

        translated_text = result['trans_result'][0]['dst']
        return translated_text

    # 有道翻译API
    APP_KEY = 'APIKEY_08fc8229'
    APP_SECRET = 'APIKEY_d3nKWPooOeWoRJXVNEX80'

    def YMT(source_text, pair):
        if pair == "中 --> En":
            source_lang = 'zh-CHS'
            target_lang = 'en'
        elif pair == "En --> 中":
            source_lang = 'en'
            target_lang = 'zh-CHS'
        else:
            raise ValueError("Invalid pair value. Expected '中 --> En' or 'En --> 中'.")

        data = {'q': source_text, 'from': source_lang, 'to': target_lang}

        addAuthParams(APP_KEY, APP_SECRET, data)

        header = {'Content-Type': 'application/x-www-form-urlencoded'}
        res = doCall('https://openapi.youdao.com/api', header, data, 'post')
        result_json = res.json()
        translated_text = result_json['translation'][0]
        return translated_text

    def doCall(url, header, params, method):
        if 'get' == method:
            return requests.get(url, params)
        elif 'post' == method:
            return requests.post(url, params, header)

    # Google翻译API
    def GMT(source_text, pair):
        url_gmt = "https://google-translator9.p.rapidapi.com/v2"

        if pair == "中 --> En":
            source_lang = 'zh'
            target_lang = 'en'
        elif pair == "En --> 中":
            source_lang = 'en'
            target_lang = 'zh'
        else:
            raise ValueError("Invalid pair value. Expected '中 --> En' or 'En --> 中'.")

        payload = {
            "q": source_text,  
            "source": source_lang,
            "target": target_lang
        }
        headers = {
            "x-rapidapi-key": "APIKEY_2782371jsnc46ecbec641c",
            "x-rapidapi-host": "google-translator9.p.rapidapi.com",
            "Content-Type": "application/json"
        }

        response = requests.post(url_gmt, json=payload, headers=headers)
        
        if response.status_code != 200:
            raise ValueError(f"Translation failed with status code {response.status_code}: {response.text}")
        
        result = response.json()
        translated_text = result['data']['translations'][0]['translatedText']
        return translated_text

    # 计算BLEU值
    def calculate_bleu_score(MT_text, refer_text, target_lang):
        def tokenize_text(text, lang):
            if lang == "en":
                return text.split()
            else:
                return list(text)
        
        refer = [tokenize_text(refer_text, target_lang)]
        target = tokenize_text(MT_text, target_lang)
        smoothing_function = SmoothingFunction().method1
        bleu_score = sentence_bleu(refer, target, smoothing_function=smoothing_function)
        return bleu_score

    # GEMBA-MQM错误的计数器函数
    def count_errors_by_severity(errors):
        severity_counts = Counter(error.severity for error in errors if error.error_type != 'no-error')
        return severity_counts

    def count_errors_by_type(errors):
    # 提取 / 前面的部分作为 error_type
        processed_error_types = (error.error_type.split('/')[0] for error in errors if error.error_type != 'no-error')
        error_type_counts = Counter(processed_error_types)
        return error_type_counts

    # 计算GEMBA-MQM得分函数
    def calculate_translation_score(text, severity_counts, lang):
        minor_weight = 5
        major_weight = 10
        critical_weight = 20
        
        total_error_score = (severity_counts.get('minor', 0) * minor_weight +
                             severity_counts.get('major', 0) * major_weight +
                             severity_counts.get('critical', 0) * critical_weight)
        
        if lang == "en":
            word_count = len(text.split())
        else:
            word_count = len(list(text))

        full_original_score = word_count

        adjusted_score = (full_original_score - total_error_score) / full_original_score * 100
        adjusted_score = round(adjusted_score, 2)

        if adjusted_score < 0:
            adjusted_score = 0

        return adjusted_score
    
    #生成MT译文、Translation Errors和PNG
    def generate_tt_and_error(MT_names, MT_texts, MQM_scores, MQM_annotations):

        # 插入各个机翻的评估结果
        for i, (MT_name, MT_text, MQM_score, errors) in enumerate(zip(MT_names, MT_texts, MQM_scores, MQM_annotations)):
            # 添加机翻标题和译文
            doc.add_heading(f'{i + 1}. {MT_name}', level=2)
            doc.add_paragraph(f'a. 译文：{MT_texts[i]}')
            doc.add_paragraph(f'b. BLEU得分：{BLEU_scores[i]}')
            doc.add_paragraph(f'c. MQM得分：{MQM_scores[i]}')
            # 添加MQM指出的翻译错误表格
            doc.add_paragraph(f'd.MQM指出的翻译错误：')
            table = doc.add_table(rows=1, cols=3)
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Severity'
            hdr_cells[1].text = 'Error Type'
            hdr_cells[2].text = 'Target'
            for error in errors:
                row_cells = table.add_row().cells
                row_cells[0].text = error['severity']
                row_cells[1].text = error['error_type']
                row_cells[2].text = error['target']

            # doc.add_picture(f'{MT_name}.png', width=Inches(4.0))
            # 添加空行分隔各个机翻结果
            doc.add_paragraph()
    
    #生成错误严重性分类
    def generate_severity_count(data):
        doc.add_heading('错误严重性分类', level=2)
        table = doc.add_table(rows=len(data) + 1, cols=4)  
        # 填充表头  
        hdr_cells = table.rows[0].cells  
        hdr_cells[0].text = 'MT Name'  
        hdr_cells[1].text = 'Minor Count'  
        hdr_cells[2].text = 'Major Count'  
        hdr_cells[3].text = 'Critical Count'    
        # 填充数据行  
        for i, item in enumerate(data, start=1):  
            row_cells = table.rows[i].cells  
            # 假设每个字典只有一个键，即MT名称  
            mt_name = list(item.keys())[0]  # 获取MT名称（注意：如果字典有多个键，这将不会工作）  
            row_cells[0].text = mt_name  
            row_cells[1].text = str(item[mt_name][0])  # Minor Count  
            row_cells[2].text = str(item[mt_name][1])  # Major Count  
            row_cells[3].text = str(item[mt_name][2])  # Critical Count  
    
    #生成错误种类分类
    def generate_error_type_count(data):
        doc.add_heading('错误种类分类', level=2)
        # 添加表格（包括表头）  
        table = doc.add_table(rows=len(data) + 1, cols=8)  
        
        # 填充表头  
        hdr_cells = table.rows[0].cells  
        hdr_cells[0].text = 'MT Name'  
        hdr_cells[1].text = 'Accuracy'  
        hdr_cells[2].text = 'Fluency'  
        hdr_cells[3].text = 'Locale Convention'  
        hdr_cells[4].text = 'Style'  
        hdr_cells[5].text = 'Terminology'  
        hdr_cells[6].text = 'Non Translation'  
        hdr_cells[7].text = 'Other'  
        
        # 填充数据行  
        for i, item in enumerate(data, start=1):  
            row_cells = table.rows[i].cells  
            # 假设每个字典只有一个键，即MT名称  
            mt_name = list(item.keys())[0]  # 获取MT名称（注意：如果字典有多个键，这将不会工作）  
            row_cells[0].text = mt_name  
            row_cells[1].text = str(item[mt_name][0])  
            row_cells[2].text = str(item[mt_name][1])
            row_cells[3].text = str(item[mt_name][2])
            row_cells[4].text = str(item[mt_name][3])
            row_cells[5].text = str(item[mt_name][4])
            row_cells[6].text = str(item[mt_name][5])
            row_cells[7].text = str(item[mt_name][6])


    os.environ["OPENAI_API_KEY"] = "APIKEY_b407cA482Ce83022c8b4c"
    os.environ["OPENAI_BASE_URL"] = "APIKEY_URL"
    llm = ChatOpenAI(model="gpt-4o")
    from langchain_core.prompts import ChatPromptTemplate

    prompt = ChatPromptTemplate.from_messages([
        ("system", """你是机器翻译质量的注解者。您的任务是使用GEMBA-MQM方法识别错误并评估翻译质量。请遵守该方法如下案例中，MQM annotations的输出方式：
        English source: Talks have resumed in Vienna to try to revive the nuclear pact, with both sides trying to gauge the prospects of success after the latest exchanges in the stop-start negotiations.
        Czech transation: Ve Vídni se ve Vídni obnovily rozhovory o oživení jaderného paktu, přičemže obě partaje se snaží posoudit vyhlídky na úspěch po posledních výměnách v jednáních.
        MQM annotations:
        critical:
        no-error
        major:
        accuracy/addition - "ve Vídni"
        accuracy/omission - "the stop-start"
        minor:
        terminology/inappropriate for context - "partaje"
        Chinese source: 大众点评乌鲁木齐家居商场频道为您提供高铁居然之家地址话，营业时间等最新商户信息，找装修公司，就上大众点评
        English translation: Urumqi Home Furnishing Store Channel provides yowith the latest business information such as the address, telephonnumber, business hours, etc., of high-speed rail, and find a decoratiocompany, and go to the reviews.
        MQM annotations:
        critical:
        accuracy/addition - "of high-speed rail"
        major:
        accuracy/mistranslation - "go to the reviews"
        minor:
        style/awkward - "etc.,"
        """),
        ("user", "{input}")
    ])

    class TranslationError(BaseModel):
        severity: str = Field(description="Each error is classified as one of three severities: critical, major, and minor. Critical errors inhibit comprehension of the text. Major errors disrupt the flow, but what the text is trying to say is still understandable. Minor errors are technically errors, but do not disrupt the flow or hinder comprehension.")
        error_type: str = Field(description="There are 7 types of error: accuracy (addition, mistranslation, omission, untranslated text), fluency (character encoding, grammar, inconsistency, punctuation, register, spelling), locale convention (currency, date, name, telephone, or time format) style (awkward), terminology (inappropriate for context, inconsistent use), non-translation, other. Please follow the output format of 'accuracy/addition'.")
        target: str = Field(description="The target place of the error.")

    class TranslationErrorReport(BaseModel):
        errors: list[TranslationError] = Field(description="List of translation errors")

    structured_llm = llm.with_structured_output(TranslationErrorReport)
    chain = prompt | structured_llm

    MT_names = ["百度翻译", "有道翻译", "谷歌翻译"]
    for MT_name in MT_names:
        if MT_name=="百度翻译":
            MT_baidu=BMT(source_text, pair)
            MT_texts.append(MT_baidu)
        if MT_name=="有道翻译":
            MT_youdao=YMT(source_text, pair)
            MT_texts.append(MT_youdao)
        if MT_name=="谷歌翻译":
            MT_google=GMT(source_text, pair)
            MT_texts.append(MT_google)
    MT_texts = [html.unescape(text) for text in MT_texts]
    return_MQM_result = []
    n = 0

    if refer_text != "":
        BLEU_scores = []
        for MT_text in MT_texts:
            BLEU_scores.append(round(calculate_bleu_score(MT_text, refer_text, target_lang), 5)*100)
        print(f"{Colors.OKGREEN}返回的BLEU分数集合{BLEU_scores}{Colors.ENDC}") 
    MQM_scores = []
    MQM_annotations = []
    error_counts = []
    severity_count_for_3_MTs = []
    error_type_count_for_3_MTs = []

    for MT_text in MT_texts:
        error_count = []
        input_data = source_lang + ": " + source_text + "; \n" + target_lang + ": " + MT_text
        print()
        print(f"{Colors.OKGREEN}--------------------{MT_names[n]}的结果如下--------------------{Colors.ENDC}")
        print()
        response = chain.invoke(input_data)
        severity_counts = count_errors_by_severity(response.errors)
        error_type_counts = count_errors_by_type(response.errors)
        print(response.errors)
        MQM_score = calculate_translation_score(MT_text, severity_counts, target_lang)
        print(f"{Colors.OKBLUE}{MT_names[n]}译文：{MT_texts[n]}{Colors.ENDC}")

        print(f"{Colors.OKGREEN}不同严重性的错误数如下：{Colors.ENDC}")
        minor_count = severity_counts.get('minor', 0)
        print("Minor errors:", minor_count)
        major_count = severity_counts.get('major', 0)
        print("Major errors:", major_count)
        critical_count = severity_counts.get('critical', 0)
        print("Critical errors:", critical_count)
        error_count.append([minor_count, major_count, critical_count])
        severity_count_for_3_MTs.append({MT_names[n]:[minor_count, major_count, critical_count]})
        

        print(f"{Colors.OKGREEN}错误具体种类数如下：{Colors.ENDC}")
        accuracy_count = error_type_counts.get('accuracy', 0)
        print("Accuracy errors:", accuracy_count)
        fluency_count = error_type_counts.get('fluency', 0)
        print("Fluency errors:", fluency_count)
        locale_convention_count = error_type_counts.get('locale convention', 0)
        print("Locale convention errors:", locale_convention_count)
        style_count = error_type_counts.get('style', 0)
        print("Style errors:", style_count)
        terminology_count = error_type_counts.get('terminology', 0)
        print("Terminology errors:", terminology_count)
        non_translation_count = error_type_counts.get('non-translation', 0)
        print("Non-translation errors:", non_translation_count)
        other_count = error_type_counts.get('other', 0)
        print("Other errors:", other_count)
        error_count.append([accuracy_count, fluency_count, locale_convention_count, style_count, terminology_count, non_translation_count, other_count])
        error_type_count_for_3_MTs.append({MT_names[n]:[accuracy_count, fluency_count, locale_convention_count, style_count, terminology_count, non_translation_count, other_count]})


        print(f"{Colors.OKGREEN}Translation Quality Score: {MQM_score}{Colors.ENDC}")
        print(response)
        MQM_annotation = [error.dict() for error in response.errors]
        MQM_scores.append(MQM_score)
        MQM_annotations.append(MQM_annotation)
        error_counts.append(error_count)
        return_MQM_result.append([MQM_annotation, MQM_score])
        n = n + 1
        print()
    generate_severity_count(severity_count_for_3_MTs)
    generate_error_type_count(error_type_count_for_3_MTs)
    print(f"{Colors.OKGREEN}--------------------MQM总结果如下--------------------{Colors.ENDC}")
    print()
    print(f"最终所有MQM相关输出见下：{return_MQM_result}")
    single_result.append([source_lang, target_lang, source_text, refer_text])
    single_result.append([MT_texts, BLEU_scores, MQM_scores, MQM_annotations, error_counts])
    generate_tt_and_error(MT_names, MT_texts, MQM_scores, MQM_annotations)
    return single_result
    

def charts(total_result):
    engines = ['百度翻译', '有道翻译', '谷歌翻译']
    error_types = ['critical', 'major', 'minor']
    type_names = ['Accuracy', 'Fluency', 'Locale convention', 'Style', 'Terminology', 'Non-translation', 'Other']

    # 初始化错误计数和类型计数字典
    error_counts = {engine: {error_type: 0 for error_type in error_types} for engine in engines}
    type_counts = {engine: {type_name: 0 for type_name in type_names} for engine in engines}
    BLEU_sum = {engine: 0 for engine in engines}
    MQM_sum = {engine: 0 for engine in engines}
    BLEU_final = {engine: 0 for engine in engines}
    MQM_final = {engine: 0 for engine in engines}
    word_count = {engine: 0 for engine in engines}

    # 遍历结果并更新错误计数和类型计数
    for item in total_result:
        for i, service in enumerate(engines):
            BLEU_sum[service] += item[1][1][i]
            MQM_sum[service] += item[1][2][i]
            word_count[service] += len(item[1][0][i])
            BLEU_final[service] += item[1][1][i] * len(item[1][0][i])
            MQM_final[service] += item[1][2][i] * len(item[1][0][i])

        for engine, index in zip(engines, range(len(engines))):
            error_type_counts = item[1][4][index][0]
            type_counts_data = item[1][4][index][1]
            
            # 更新错误计数
            for error_type, count in zip(error_types, error_type_counts):
                error_counts[engine][error_type] += count
            
            # 更新类型计数
            for type_name, count in zip(type_names, type_counts_data):
                type_counts[engine][type_name] += count

    # 计算得分
    BLEU_weigh = {engine: BLEU_final[engine] / word_count[engine] for engine in engines}
    MQM_weigh = {engine: MQM_final[engine] / word_count[engine] for engine in engines}

    # 打印错误计数结果
    y1_arrays = {}
    for engine in error_counts:
        y1_arrays[engine] = np.array([error_counts[engine]['critical'], error_counts[engine]['major'], error_counts[engine]['minor']])

    # 打印类型计数结果
    y2_arrays = {}
    for engine in type_counts:
        y2_arrays[engine] = np.array([type_counts[engine]['Accuracy'], type_counts[engine]['Fluency'], type_counts[engine]['Locale convention'], type_counts[engine]['Style'], type_counts[engine]['Terminology'], type_counts[engine]['Non-translation'], type_counts[engine]['Other']])

    # 设置matplotlib支持中文显示
    plt.rcParams['font.sans-serif'] = ['Microsoft YaHei'] 
    plt.rcParams['font.family']='sans-serif'
    plt.rcParams['axes.unicode_minus'] = False

    # 过滤掉0值对应的数据点和标签
    def filter_data_labels(data, labels):
        filtered_data = []
        filtered_labels = []
        for value, label in zip(data, labels):
            if value != 0:
                filtered_data.append(value)
                filtered_labels.append(label)
        return np.array(filtered_data), filtered_labels

    # 绘制错误严重性饼图
    def plot_pie_charts(engine, y1_data, y1_labels, y2_data, y2_labels):
        fig, axs = plt.subplots(1, 2, figsize=(12, 6))
        
        axs[0].pie(y1_data, labels=y1_labels, autopct='%.2f%%', colors=["#f89588", "#63b2ee", "#76da91"])
        axs[0].set_title("错误严重性")
        
        axs[1].pie(y2_data, labels=y2_labels, autopct='%.2f%%', colors=["#95a2ff", "#fa8080", "#ffc076", '#87e885', '#3cb9fc', '#cb9bff', '#9987ce'])
        axs[1].set_title("错误类型")
        
        plt.suptitle(f'{engine}翻译', fontsize=16)
        plt.savefig(f"{engine}.png")
        plt.close(fig)

    # 绘制得分柱状图
    def plot_score_charts():
        bar_names = ["百度翻译", "有道翻译", "Google翻译"]
        fig, axs = plt.subplots(1, 3, figsize=(18, 6))

        # 绘制MQM得分情况
        axs[0].bar(bar_names, [MQM_weigh[engine] for engine in engines], color='#FF7F3E', label='MQM')
        axs[0].set_title('MQM得分情况')
        axs[0].set_ylim(min(MQM_weigh.values()) * 0.9, max(MQM_weigh.values()) * 1.1)
        axs[0].legend()
        for i, v in enumerate([MQM_weigh[engine] for engine in engines]):
            axs[0].text(i, v + 0.01, f"{v:.2f}", ha='center', va='bottom')

        # 绘制BLEU得分情况
        axs[1].bar(bar_names, [BLEU_weigh[engine] for engine in engines], color='#2A629A', label='BLEU')
        axs[1].set_title('BLEU得分情况')
        axs[1].set_ylim(min(BLEU_weigh.values()) * 0.9, max(BLEU_weigh.values()) * 1.1)
        axs[1].legend()
        for i, v in enumerate([BLEU_weigh[engine] for engine in engines]):
            axs[1].text(i, v + 0.01, f"{v:.2f}", ha='center', va='bottom')

        # 绘制MQM与BLEU得分情况对比
        bar_width = 0.35
        index = range(len(bar_names))
        bar1 = axs[2].bar(index, [MQM_weigh[engine] for engine in engines], bar_width, label='MQM', color='#FF7F3E')
        bar2 = axs[2].bar([i + bar_width for i in index], [BLEU_weigh[engine] for engine in engines], bar_width, label='BLEU', color='#2A629A')
        axs[2].set_title('MQM与BLEU得分情况对比')
        axs[2].set_xticks([i + bar_width / 2 for i in index])
        axs[2].set_xticklabels(bar_names)
        axs[2].legend()
        for i, v1, v2 in zip(index, [MQM_weigh[engine] for engine in engines], [BLEU_weigh[engine] for engine in engines]):
            axs[2].text(i, v1 + 0.01, f"{v1:.2f}", ha='center', va='bottom')
            axs[2].text(i + bar_width, v2 + 0.01, f"{v2:.2f}", ha='center', va='bottom')

        plt.savefig('score_chart.png')

    # 绘制每个引擎的错误类型和得分图表
    for engine in engines:
        y1_data, y1_labels = filter_data_labels(y1_arrays[engine], ['Critical', 'Major', 'Minor'])
        y2_data, y2_labels = filter_data_labels(y2_arrays[engine], ['Accuracy', 'Fluency', 'Locale convention', 'Style', 'Terminology', 'Non-translation', 'Other'])
        plot_pie_charts(engine, y1_data, y1_labels, y2_data, y2_labels)

    # 绘制得分图表
    plot_score_charts()


def generate_pair_st_rt(pair,source_text,refer_text):
        doc.add_paragraph(f'语言对：{pair}')
        doc.add_paragraph(f'Source Text：{source_text}')
        doc.add_paragraph(f'Refer Text：{refer_text}')

def generate_score_chart():
    # 添加标题与说明
    doc.add_heading('最终返回报告', level=1)
    doc.add_paragraph(f'该次翻译实践的语言对：{pair}')
    doc.add_paragraph(f'说明：下文报告对以上所有运行结果进行汇总得出，最终MQM得分是计算了基于字数的加权平均数。')
    doc.add_paragraph(f'MQM得分原理介绍：非GEMBA-MQM官方，但基于其结果，将critical，major，和minor分别以20分，10分，5分的标准来计算总错误扣分，然后基于每100词总原始分加100分来计算，最后用（总原始分-总错误分/总原始分）*100作为GEMBA-MQM的量化分数。')
    # 插入图片
    doc.add_paragraph(f'得分对比')
    doc.add_picture('score_chart.png', width=Inches(6))  # 调整宽度适合页面
    doc.add_paragraph(f'各机器翻译错误类型占比')
    MT_names = ["百度翻译","有道翻译","谷歌翻译"]
    for MT_name in MT_names:
        doc.add_picture(f'{MT_name}.png', width=Inches(4.0))

def main(seg_number, pair):
    print(f"将连续接受{seg_number}次文本")
    for times in range(0, int(seg_number)):
        print()
        print(f"{Colors.OKGREEN}--------------------第{times+1}次运行中--------------------{Colors.ENDC}")
        doc.add_heading(f'第{times+1}片段', level=1)
        get_user_input(pair)
        generate_pair_st_rt(pair,source_text,refer_text)
        single_result = run_translation(refer_text, source_text, pair)
        total_result.append(single_result)
    charts(total_result)
    print(total_result)
    generate_score_chart()
    doc.save('translation_quality_evaluation_report.docx')
    return total_result

if __name__ == "__main__":
    preparation()
